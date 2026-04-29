from docx import Document
import json

def extract_docx_content(docx_path):
    doc = Document(docx_path)
    content = []
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append({'type': 'paragraph', 'text': text})
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content.append({'type': 'table', 'data': table_data})
    # Extract section titles (if styled)
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            content.append({'type': 'heading', 'text': para.text.strip()})
    return content

if __name__ == '__main__':
    docx_path = 'BUS100 - 01_E4E _23510-Schedule_SPR 2026-2.docx'
    content = extract_docx_content(docx_path)
    print(json.dumps(content, indent=2, ensure_ascii=False))
